import logging

import numpy as np
from docx import Document as DocxDocument
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

RERANK_THRESHOLD = 0.3  # Normalized score threshold (0.3 ≈ -0.4 raw cosine sim)


def load_documents(uploaded_files):
    docs = []
    for f in uploaded_files:
        try:
            if f.name.endswith(".pdf"):
                reader = PdfReader(f)
                text = "".join(p.extract_text() or "" for p in reader.pages)
                docs.append(Document(page_content=text, metadata={"source": f.name}))
            elif f.name.endswith(".docx"):
                docx = DocxDocument(f)
                text = "\n".join(p.text for p in docx.paragraphs)
                docs.append(Document(page_content=text, metadata={"source": f.name}))
        except Exception:
            logging.exception("Error reading file: %s", f.name)
    return docs


def build_retriever(uploaded_files, k: int = 6):
    docs = load_documents(uploaded_files)
    if not docs:
        return None
    splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=80)
    chunks = splitter.split_documents(docs)
    if not chunks:
        return None
    try:
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        vectorstore = FAISS.from_documents(chunks, embeddings)
        # Use plain similarity search (not similarity_search_with_relevance_scores)
        # to avoid the unnormalized score warning. All scoring is done in rerank().
        return vectorstore.as_retriever(
            search_type="similarity", search_kwargs={"k": k}
        )
    except Exception:
        logging.exception("Failed to build retriever.")
        return None


def rerank(query: str, docs: list) -> list:
    """
    Score every retrieved chunk against the query using cosine similarity.
    Drop chunks below RERANK_THRESHOLD so genuinely irrelevant chunks
    never reach the LLM — but only after we already have candidates.
    Scores are normalized to [0, 1] range for LangChain compatibility.
    """
    if not docs:
        return docs
    try:
        from models.embeddings import get_embedding_model

        model = get_embedding_model()
        q_emb = model.encode(query, convert_to_numpy=True)
        scored = []
        for doc in docs:
            d_emb = model.encode(doc.page_content, convert_to_numpy=True)
            # Compute cosine similarity (ranges from -1 to 1)
            cosine_sim = float(
                np.dot(q_emb, d_emb)
                / (np.linalg.norm(q_emb) * np.linalg.norm(d_emb) + 1e-9)
            )
            # Normalize to [0, 1] range for LangChain compatibility
            normalized_score = (cosine_sim + 1.0) / 2.0
            if normalized_score >= RERANK_THRESHOLD:
                scored.append((normalized_score, doc))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in scored]
    except Exception:
        logging.exception("Reranking failed, returning original order.")
        return docs


def get_context(retriever, query: str) -> tuple:
    """
    Retrieve -> rerank with threshold filter -> return top chunks as context.
    Returns (context_text, source_label). Both empty if nothing passes threshold.
    """
    if retriever is None:
        return "", ""
    try:
        docs = retriever.invoke(query)
        if not docs:
            return "", ""
        docs = rerank(query, docs)
        if not docs:
            return "", ""
        sources = list({d.metadata.get("source", "document") for d in docs})
        context = "\n\n".join(d.page_content for d in docs[:3])
        return context, ", ".join(sources)
    except Exception:
        logging.exception("Context retrieval failed.")
        return "", ""
